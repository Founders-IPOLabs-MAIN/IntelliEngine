"""
DRHP Document Export Module
Exports DRHP content to Word (.docx) and PDF formats while preserving SEBI-specific formatting.
"""

import io
import re
import logging
from typing import Optional, Dict, List, Tuple, Any
from docx import Document
from docx.shared import Pt, Inches, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
import weasyprint

logger = logging.getLogger(__name__)

# SEBI DRHP CSS class to Word style mapping
SEBI_CLASS_TO_STYLE = {
    'drhp-title': {'style': 'Title', 'alignment': WD_ALIGN_PARAGRAPH.CENTER, 'bold': True, 'size': 18},
    'drhp-h1': {'style': 'Heading 1', 'bold': True, 'size': 16},
    'drhp-h2': {'style': 'Heading 2', 'bold': True, 'size': 14},
    'drhp-h3': {'style': 'Heading 3', 'bold': True, 'size': 13},
    'drhp-h4': {'style': 'Heading 4', 'bold': True, 'size': 12},
    'drhp-h5': {'style': 'Heading 5', 'bold': True, 'size': 11},
    'drhp-h6': {'style': 'Heading 6', 'bold': True, 'size': 10},
    'drhp-section': {'style': 'Heading 2', 'bold': True, 'size': 14},
    'drhp-subsection': {'style': 'Heading 3', 'bold': True, 'size': 13},
    'drhp-clause': {'style': 'Normal', 'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY},
    'drhp-sub-clause': {'style': 'Normal', 'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY, 'left_indent': 0.5},
    'drhp-para': {'style': 'Normal', 'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY},
    'drhp-toc-heading': {'style': 'TOC Heading', 'bold': True, 'size': 14},
    'drhp-toc-1': {'style': 'Normal', 'left_indent': 0},
    'drhp-toc-2': {'style': 'Normal', 'left_indent': 0.25},
    'drhp-toc-3': {'style': 'Normal', 'left_indent': 0.5},
    'drhp-toc-4': {'style': 'Normal', 'left_indent': 0.75},
    'drhp-legal': {'style': 'Normal', 'italic': True, 'left_indent': 0.25},
    'drhp-disclaimer': {'style': 'Normal', 'size': 10, 'left_indent': 0.25},
    'drhp-risk-factor': {'style': 'Normal', 'left_indent': 0.25},
    'drhp-material-contract': {'style': 'Normal'},
    'drhp-list-para': {'style': 'Normal'},
    'drhp-list-number': {'style': 'Normal'},
    'drhp-list-bullet': {'style': 'Normal'},
    'drhp-quote': {'style': 'Quote', 'italic': True, 'left_indent': 0.5},
    'drhp-quote-intense': {'style': 'Intense Quote', 'italic': True, 'bold': True, 'left_indent': 0.5},
    'drhp-normal': {'style': 'Normal'},
    'drhp-body': {'style': 'Normal', 'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY},
    'drhp-body-2': {'style': 'Normal', 'left_indent': 0.25},
    'drhp-body-3': {'style': 'Normal', 'left_indent': 0.5},
}

# Alignment CSS to Word mapping
CSS_ALIGNMENT_MAP = {
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class DRHPWordExporter:
    """
    Export DRHP HTML content to Word document with SEBI formatting preservation.
    """
    
    def __init__(self, html_content: str, company_name: str = "Company", board_type: str = "sme"):
        self.html_content = html_content
        self.company_name = company_name
        self.board_type = board_type
        self.doc = None
        self.current_section = None
        self.list_counter = 0
        self.warnings: List[str] = []
        
    def export(self) -> bytes:
        """
        Export HTML content to Word document.
        Returns the document as bytes.
        """
        try:
            # Create a new document
            self.doc = Document()
            
            # Set up document properties
            self._setup_document_properties()
            
            # Set up custom styles
            self._setup_custom_styles()
            
            # Parse HTML and convert to Word
            soup = BeautifulSoup(self.html_content, 'html.parser')
            
            # Find the main content container
            content_div = soup.find('div', class_='drhp-document')
            if content_div:
                self._process_elements(content_div.children)
            else:
                # Process entire body
                self._process_elements(soup.children)
            
            # Save to bytes
            buffer = io.BytesIO()
            self.doc.save(buffer)
            buffer.seek(0)
            
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error exporting to Word: {str(e)}")
            raise
    
    def _setup_document_properties(self):
        """Set up document properties and page layout."""
        # Set page margins (SEBI DRHP standard: 1 inch margins)
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1)
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        
        # Set document core properties
        core_props = self.doc.core_properties
        core_props.title = f"DRHP - {self.company_name}"
        core_props.author = "IntelliEngine DRHP Builder"
        core_props.subject = f"{self.board_type.upper()} Board Draft Red Herring Prospectus"
    
    def _setup_custom_styles(self):
        """Set up custom Word styles for SEBI formatting."""
        styles = self.doc.styles
        
        # Ensure base styles exist and configure them
        try:
            # Configure Normal style
            normal_style = styles['Normal']
            normal_style.font.name = 'Times New Roman'
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.line_spacing = 1.5
            normal_style.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass
        
        # Try to configure heading styles
        for level in range(1, 7):
            try:
                heading_style = styles[f'Heading {level}']
                heading_style.font.name = 'Times New Roman'
                heading_style.font.bold = True
                heading_style.font.size = Pt(18 - level * 2) if level <= 3 else Pt(12)
            except KeyError:
                pass
    
    def _process_elements(self, elements):
        """Process HTML elements and convert to Word."""
        for element in elements:
            if element.name is None:
                # Text node
                text = str(element).strip()
                if text:
                    self._add_text_run(text)
            elif element.name == 'style':
                # Skip style tags
                continue
            elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                self._process_heading(element)
            elif element.name == 'p':
                self._process_paragraph(element)
            elif element.name == 'table':
                self._process_table(element)
            elif element.name in ['ul', 'ol']:
                self._process_list(element)
            elif element.name == 'blockquote':
                self._process_blockquote(element)
            elif element.name == 'hr':
                self._process_page_break(element)
            elif element.name == 'div':
                # Process children of div
                self._process_elements(element.children)
            elif element.name == 'br':
                # Line break
                if self.current_paragraph:
                    self.current_paragraph.add_run('\n')
            elif element.name == 'img':
                self._process_image(element)
    
    def _process_heading(self, element):
        """Process heading elements (h1-h6)."""
        level = int(element.name[1])
        text = element.get_text()
        
        # Get CSS classes
        classes = element.get('class', [])
        style_info = self._get_style_info(classes)
        
        # Add heading paragraph
        para = self.doc.add_paragraph()
        
        # Try to apply Word heading style
        try:
            para.style = f'Heading {level}'
        except KeyError:
            pass
        
        # Apply formatting from style info
        self._apply_paragraph_formatting(para, element, style_info)
        
        # Add text with formatting
        run = para.add_run(text)
        run.bold = True
        if 'size' in style_info:
            run.font.size = Pt(style_info['size'])
        run.font.name = 'Times New Roman'
    
    def _process_paragraph(self, element):
        """Process paragraph elements."""
        # Get CSS classes and inline styles
        classes = element.get('class', [])
        style_info = self._get_style_info(classes)
        
        # Add paragraph
        para = self.doc.add_paragraph()
        self.current_paragraph = para
        
        # Apply formatting
        self._apply_paragraph_formatting(para, element, style_info)
        
        # Process inline content
        self._process_inline_content(element, para)
    
    def _process_inline_content(self, element, para):
        """Process inline content (text, bold, italic, links, etc.)."""
        for child in element.children:
            if child.name is None:
                # Text node
                text = str(child)
                if text:
                    para.add_run(text)
            elif child.name == 'strong' or child.name == 'b':
                run = para.add_run(child.get_text())
                run.bold = True
            elif child.name == 'em' or child.name == 'i':
                run = para.add_run(child.get_text())
                run.italic = True
            elif child.name == 'u':
                run = para.add_run(child.get_text())
                run.underline = True
            elif child.name == 's' or child.name == 'strike':
                run = para.add_run(child.get_text())
                run.font.strike = True
            elif child.name == 'sup':
                run = para.add_run(child.get_text())
                run.font.superscript = True
            elif child.name == 'sub':
                run = para.add_run(child.get_text())
                run.font.subscript = True
            elif child.name == 'span':
                self._process_span(child, para)
            elif child.name == 'a':
                self._process_link(child, para)
            elif child.name == 'br':
                para.add_run('\n')
            elif child.name == 'mark':
                # Highlighted text
                run = para.add_run(child.get_text())
                run.font.highlight_color = 7  # Yellow
            else:
                # Recursively process other inline elements
                self._process_inline_content(child, para)
    
    def _process_span(self, element, para):
        """Process span elements with inline styles."""
        text = element.get_text()
        run = para.add_run(text)
        
        # Parse inline styles
        style = element.get('style', '')
        styles = self._parse_inline_style(style)
        
        # Apply font styles
        if 'font-size' in styles:
            size = self._parse_size(styles['font-size'])
            if size:
                run.font.size = Pt(size)
        
        if 'color' in styles:
            color = self._parse_color(styles['color'])
            if color:
                run.font.color.rgb = color
        
        if 'background-color' in styles:
            # Highlight - approximate with yellow
            run.font.highlight_color = 7
        
        if 'font-weight' in styles and styles['font-weight'] == 'bold':
            run.bold = True
        
        if 'font-style' in styles and styles['font-style'] == 'italic':
            run.italic = True
        
        if 'text-decoration' in styles and 'underline' in styles['text-decoration']:
            run.underline = True
    
    def _process_link(self, element, para):
        """Process hyperlink elements."""
        text = element.get_text()
        # href = element.get('href', '')  # Reserved for future hyperlink implementation
        
        # Add as blue underlined text
        run = para.add_run(text)
        run.font.color.rgb = RGBColor(0x1D, 0xA1, 0xF2)  # Blue
        run.underline = True
        
        # Add actual hyperlink (requires more complex XML manipulation)
        # For now, just style it as a link
    
    def _process_table(self, element):
        """Process table elements."""
        rows = element.find_all('tr')
        if not rows:
            return
        
        # Count columns
        first_row = rows[0]
        cols = len(first_row.find_all(['td', 'th']))
        
        if cols == 0:
            return
        
        # Create table
        table = self.doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for col_idx, cell in enumerate(cells):
                if col_idx < cols:
                    table_cell = table.rows[row_idx].cells[col_idx]
                    
                    # Clear default paragraph and add content
                    if table_cell.paragraphs:
                        para = table_cell.paragraphs[0]
                    else:
                        para = table_cell.add_paragraph()
                    
                    # Process cell content
                    self._process_inline_content(cell, para)
                    
                    # Bold for header cells
                    if cell.name == 'th':
                        for run in para.runs:
                            run.bold = True
        
        # Add spacing after table
        self.doc.add_paragraph()
    
    def _process_list(self, element, level=0):
        """Process list elements (ul/ol)."""
        is_ordered = element.name == 'ol'
        counter = 1
        
        for item in element.find_all('li', recursive=False):
            # Create paragraph with appropriate indentation
            para = self.doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
            
            # Add bullet or number
            if is_ordered:
                prefix = f"{counter}.\t"
                counter += 1
            else:
                prefix = "•\t"
            
            para.add_run(prefix)
            
            # Process item content
            self._process_inline_content(item, para)
            
            # Check for nested lists
            nested_list = item.find(['ul', 'ol'])
            if nested_list:
                self._process_list(nested_list, level + 1)
    
    def _process_blockquote(self, element):
        """Process blockquote elements."""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.right_indent = Inches(0.5)
        
        # Process content
        self._process_inline_content(element, para)
        
        # Apply italic formatting
        for run in para.runs:
            run.italic = True
    
    def _process_page_break(self, element):
        """Process page break (hr with page-break class)."""
        classes = element.get('class', [])
        if 'drhp-page-break' in classes or 'page-break' in classes:
            self.doc.add_page_break()
        else:
            # Regular horizontal rule - add a paragraph with border
            para = self.doc.add_paragraph()
            para.add_run('─' * 50)
    
    def _process_image(self, element):
        """Process image elements."""
        src = element.get('src', '')
        
        if src.startswith('data:'):
            # Base64 image - skip for now (would need to decode and add)
            self.warnings.append("Embedded image skipped - export as separate file")
        elif src.startswith('/api/'):
            # API-served image - would need to fetch
            self.warnings.append(f"Image reference: {src}")
        
        # Add placeholder text
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run('[Image]')
        run.italic = True
    
    def _get_style_info(self, classes) -> Dict:
        """Get style information from CSS classes."""
        style_info = {}
        
        if isinstance(classes, str):
            classes = classes.split()
        
        for cls in classes:
            if cls in SEBI_CLASS_TO_STYLE:
                style_info.update(SEBI_CLASS_TO_STYLE[cls])
        
        return style_info
    
    def _apply_paragraph_formatting(self, para, element, style_info: Dict):
        """Apply paragraph formatting based on element and style info."""
        pf = para.paragraph_format
        
        # Alignment from style info
        if 'alignment' in style_info:
            pf.alignment = style_info['alignment']
        
        # Alignment from inline style
        inline_style = element.get('style', '')
        styles = self._parse_inline_style(inline_style)
        
        if 'text-align' in styles:
            align = styles['text-align']
            if align in CSS_ALIGNMENT_MAP:
                pf.alignment = CSS_ALIGNMENT_MAP[align]
        
        # Left indent
        if 'margin-left' in styles:
            indent = self._parse_size(styles['margin-left'])
            if indent:
                pf.left_indent = Pt(indent)
        elif 'left_indent' in style_info:
            pf.left_indent = Inches(style_info['left_indent'])
        
        # Right indent
        if 'margin-right' in styles:
            indent = self._parse_size(styles['margin-right'])
            if indent:
                pf.right_indent = Pt(indent)
        
        # First line indent
        if 'text-indent' in styles:
            indent = self._parse_size(styles['text-indent'])
            if indent:
                pf.first_line_indent = Pt(indent)
        
        # Spacing before
        if 'margin-top' in styles:
            space = self._parse_size(styles['margin-top'])
            if space:
                pf.space_before = Pt(space)
        
        # Spacing after
        if 'margin-bottom' in styles:
            space = self._parse_size(styles['margin-bottom'])
            if space:
                pf.space_after = Pt(space)
        
        # Line height
        if 'line-height' in styles:
            lh = styles['line-height']
            try:
                if lh.endswith('px'):
                    pf.line_spacing = Pt(float(lh[:-2]))
                else:
                    pf.line_spacing = float(lh)
            except (ValueError, TypeError):
                pass
    
    def _parse_inline_style(self, style_str: str) -> Dict:
        """Parse inline CSS style string."""
        styles = {}
        if not style_str:
            return styles
        
        for item in style_str.split(';'):
            item = item.strip()
            if ':' in item:
                key, value = item.split(':', 1)
                styles[key.strip()] = value.strip()
        
        return styles
    
    def _parse_size(self, size_str: str) -> Optional[float]:
        """Parse CSS size value to points."""
        if not size_str:
            return None
        
        size_str = size_str.strip()
        
        try:
            if size_str.endswith('pt'):
                return float(size_str[:-2])
            elif size_str.endswith('px'):
                return float(size_str[:-2]) * 0.75  # px to pt
            elif size_str.endswith('em'):
                return float(size_str[:-2]) * 12  # Assuming 12pt base
            elif size_str.endswith('rem'):
                return float(size_str[:-3]) * 12
            elif size_str.endswith('%'):
                return float(size_str[:-1]) * 0.12  # 100% = 12pt
            else:
                return float(size_str)
        except (ValueError, TypeError):
            return None
    
    def _parse_color(self, color_str: str) -> Optional[RGBColor]:
        """Parse CSS color value to RGBColor."""
        if not color_str:
            return None
        
        color_str = color_str.strip()
        
        try:
            if color_str.startswith('#'):
                # Hex color
                hex_color = color_str[1:]
                if len(hex_color) == 3:
                    hex_color = ''.join([c*2 for c in hex_color])
                if len(hex_color) == 6:
                    r = int(hex_color[0:2], 16)
                    g = int(hex_color[2:4], 16)
                    b = int(hex_color[4:6], 16)
                    return RGBColor(r, g, b)
            elif color_str.startswith('rgb'):
                # rgb(r, g, b) format
                match = re.search(r'rgb\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', color_str)
                if match:
                    r, g, b = map(int, match.groups())
                    return RGBColor(r, g, b)
        except (ValueError, TypeError):
            pass
        
        return None
    
    def _add_text_run(self, text: str):
        """Add a text run to the current context."""
        # This is for standalone text nodes
        if not text.strip():
            return
        
        para = self.doc.add_paragraph()
        para.add_run(text)


class DRHPPDFExporter:
    """
    Export DRHP HTML content to PDF with SEBI formatting preservation.
    Uses WeasyPrint for high-fidelity HTML to PDF conversion.
    """
    
    def __init__(self, html_content: str, company_name: str = "Company", board_type: str = "sme"):
        self.html_content = html_content
        self.company_name = company_name
        self.board_type = board_type
    
    def export(self) -> bytes:
        """
        Export HTML content to PDF.
        Returns the PDF as bytes.
        """
        try:
            # Wrap content in full HTML document with print styles
            full_html = self._create_print_document()
            
            # Generate PDF using WeasyPrint
            pdf_bytes = weasyprint.HTML(string=full_html).write_pdf()
            
            return pdf_bytes
            
        except Exception as e:
            logger.error(f"Error exporting to PDF: {str(e)}")
            raise
    
    def _create_print_document(self) -> str:
        """Create a full HTML document with print-optimized styles."""
        return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>DRHP - {self.company_name}</title>
    <style>
        @page {{
            size: A4;
            margin: 2.5cm 2cm 2.5cm 2.5cm;
            @top-center {{
                content: "DRAFT RED HERRING PROSPECTUS - {self.company_name.upper()}";
                font-family: 'Times New Roman', Times, serif;
                font-size: 9pt;
                color: #666;
            }}
            @bottom-center {{
                content: counter(page);
                font-family: 'Times New Roman', Times, serif;
                font-size: 10pt;
            }}
        }}
        
        @page :first {{
            @top-center {{
                content: "";
            }}
        }}
        
        body {{
            font-family: 'Times New Roman', Times, serif;
            font-size: 12pt;
            line-height: 1.5;
            color: #000;
            text-align: justify;
        }}
        
        /* SEBI DRHP Specific Styles */
        .drhp-document {{
            max-width: none;
        }}
        
        .drhp-title {{
            font-size: 18pt;
            font-weight: bold;
            text-align: center;
            margin: 24px 0 16px 0;
            text-transform: uppercase;
            page-break-after: avoid;
        }}
        
        .drhp-h1, .drhp-section {{
            font-size: 16pt;
            font-weight: bold;
            margin: 20px 0 12px 0;
            page-break-after: avoid;
        }}
        
        .drhp-h2, .drhp-subsection {{
            font-size: 14pt;
            font-weight: bold;
            margin: 16px 0 10px 0;
            page-break-after: avoid;
        }}
        
        .drhp-h3 {{ font-size: 13pt; font-weight: bold; margin: 14px 0 8px 0; page-break-after: avoid; }}
        .drhp-h4 {{ font-size: 12pt; font-weight: bold; margin: 12px 0 6px 0; }}
        .drhp-h5 {{ font-size: 11pt; font-weight: bold; margin: 10px 0 4px 0; }}
        .drhp-h6 {{ font-size: 10pt; font-weight: bold; margin: 8px 0 4px 0; }}
        
        h1 {{ font-size: 18pt; font-weight: bold; margin: 24px 0 16px 0; page-break-after: avoid; }}
        h2 {{ font-size: 16pt; font-weight: bold; margin: 20px 0 12px 0; page-break-after: avoid; }}
        h3 {{ font-size: 14pt; font-weight: bold; margin: 16px 0 10px 0; page-break-after: avoid; }}
        h4 {{ font-size: 12pt; font-weight: bold; margin: 12px 0 8px 0; }}
        h5 {{ font-size: 11pt; font-weight: bold; margin: 10px 0 6px 0; }}
        h6 {{ font-size: 10pt; font-weight: bold; margin: 8px 0 4px 0; }}
        
        .drhp-clause {{
            margin: 8px 0;
            text-align: justify;
        }}
        
        .drhp-sub-clause {{
            margin: 6px 0 6px 36px;
            text-align: justify;
        }}
        
        .drhp-para {{
            margin: 8px 0;
            text-align: justify;
        }}
        
        .drhp-legal {{
            font-style: italic;
            margin: 12px 0;
            padding: 10px;
            border-left: 3px solid #333;
            background: #f9f9f9;
        }}
        
        .drhp-disclaimer {{
            font-size: 10pt;
            color: #444;
            margin: 10px 0;
            padding: 8px;
            border: 1px solid #ccc;
            background: #fafafa;
        }}
        
        .drhp-risk-factor {{
            margin: 10px 0 10px 20px;
            padding-left: 10px;
            border-left: 2px solid #d00;
        }}
        
        /* TOC Styles */
        .drhp-toc-heading {{
            font-size: 14pt;
            font-weight: bold;
            margin: 20px 0 12px 0;
            text-transform: uppercase;
        }}
        .drhp-toc-1 {{ margin: 4px 0 4px 0; }}
        .drhp-toc-2 {{ margin: 2px 0 2px 24px; }}
        .drhp-toc-3 {{ margin: 2px 0 2px 48px; }}
        .drhp-toc-4 {{ margin: 2px 0 2px 72px; }}
        
        /* Table Styles */
        table, .drhp-table {{
            width: 100%;
            border-collapse: collapse;
            margin: 16px 0;
            font-size: 11pt;
            page-break-inside: avoid;
        }}
        
        table th, table td,
        .drhp-table th, .drhp-table td {{
            border: 1px solid #000;
            padding: 8px 10px;
            vertical-align: top;
            text-align: left;
        }}
        
        table th, .drhp-table th {{
            background-color: #f0f0f0;
            font-weight: bold;
        }}
        
        /* List Styles */
        .drhp-numbered-list {{
            list-style: none;
            padding-left: 0;
            margin: 8px 0;
            counter-reset: item;
        }}
        
        .drhp-numbered-list > li {{
            counter-increment: item;
            margin: 6px 0;
            padding-left: 40px;
            position: relative;
        }}
        
        .drhp-numbered-list > li::before {{
            content: counter(item) ".";
            position: absolute;
            left: 0;
        }}
        
        .drhp-numbered-list.level-2 > li::before {{ content: "(" counter(item, lower-alpha) ")"; }}
        .drhp-numbered-list.level-3 > li::before {{ content: "(" counter(item, lower-roman) ")"; }}
        
        .drhp-bullet-list {{
            list-style: disc;
            margin: 8px 0 8px 40px;
        }}
        
        /* Quote Styles */
        .drhp-quote, blockquote {{
            margin: 16px 40px;
            padding: 12px 20px;
            border-left: 4px solid #666;
            font-style: italic;
            background: #f9f9f9;
        }}
        
        .drhp-quote-intense {{
            margin: 16px 40px;
            padding: 12px 20px;
            border-left: 4px solid #1DA1F2;
            font-style: italic;
            background: #f0f9ff;
            font-weight: 500;
        }}
        
        /* Link Styles */
        a, .drhp-link {{
            color: #1DA1F2;
            text-decoration: underline;
        }}
        
        /* Page Break */
        .drhp-page-break {{
            page-break-before: always;
            border: none;
            margin: 0;
            padding: 0;
            height: 0;
        }}
        
        hr {{
            border: none;
            border-top: 1px solid #000;
            margin: 20px 0;
        }}
        
        /* Image Styles */
        img, .drhp-image {{
            max-width: 100%;
            height: auto;
            margin: 12px 0;
        }}
        
        /* Body text */
        .drhp-normal, .drhp-body {{
            margin: 8px 0;
            text-align: justify;
        }}
        
        .drhp-body-2 {{ margin: 6px 0 6px 24px; }}
        .drhp-body-3 {{ margin: 6px 0 6px 48px; }}
        
        /* Ensure proper paragraph handling */
        p {{
            margin: 8px 0;
            text-align: justify;
            orphans: 3;
            widows: 3;
        }}
        
        /* Print-specific overrides */
        @media print {{
            body {{
                print-color-adjust: exact;
                -webkit-print-color-adjust: exact;
            }}
        }}
    </style>
</head>
<body>
    {self.html_content}
</body>
</html>'''
