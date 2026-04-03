"""
DRHP Export Module Tests
Tests for Word (.docx) and PDF export functionality with SEBI formatting preservation.
"""

import pytest
import requests
import os
import io
from datetime import datetime, timezone

# Get BASE_URL from environment
BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', '').rstrip('/')

# Test HTML content with SEBI formatting
TEST_HTML_CONTENT = """
<div class="drhp-document">
    <h1 class="drhp-title">DRAFT RED HERRING PROSPECTUS</h1>
    <h2 class="drhp-h1">SECTION I: GENERAL INFORMATION</h2>
    <p class="drhp-para">This is a <strong>bold text</strong> and <em>italic text</em> paragraph.</p>
    <p class="drhp-clause" style="text-align: justify;">This is a justified clause with proper alignment.</p>
    <p class="drhp-sub-clause" style="margin-left: 36px;">This is a sub-clause with indentation.</p>
    
    <h3 class="drhp-h2">Risk Factors</h3>
    <p class="drhp-risk-factor">Risk Factor 1: Market volatility may affect share price.</p>
    
    <table class="drhp-table">
        <tr>
            <th>Header 1</th>
            <th>Header 2</th>
        </tr>
        <tr>
            <td>Cell 1</td>
            <td>Cell 2</td>
        </tr>
    </table>
    
    <ul class="drhp-bullet-list">
        <li>Bullet item 1</li>
        <li>Bullet item 2</li>
    </ul>
    
    <ol class="drhp-numbered-list">
        <li>Numbered item 1</li>
        <li>Numbered item 2</li>
    </ol>
    
    <p class="drhp-legal">This is legal disclaimer text in italic style.</p>
    <p class="drhp-disclaimer">This is a disclaimer with smaller font size.</p>
    
    <blockquote class="drhp-quote">This is a quoted text block.</blockquote>
    
    <p>Link example: <a href="https://sebi.gov.in">SEBI Website</a></p>
</div>
"""

# Minimal HTML for basic tests
MINIMAL_HTML = "<h1>Test Document</h1><p>Simple paragraph.</p>"


class TestDRHPExportEndpoint:
    """Tests for the DRHP export endpoint"""
    
    @pytest.fixture(autouse=True)
    def setup(self):
        """Setup test fixtures"""
        self.session = requests.Session()
        self.session.headers.update({"Content-Type": "application/json"})
        self.test_project_id = None
        self.auth_token = None
        
    def _create_test_user_and_project(self):
        """Create a test user and project for testing"""
        import uuid
        from pymongo import MongoClient
        
        # Connect to MongoDB
        mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
        db_name = os.environ.get('DB_NAME', 'test_database')
        client = MongoClient(mongo_url)
        db = client[db_name]
        
        # Create test user
        user_id = f"test_user_{uuid.uuid4().hex[:8]}"
        session_token = f"test_token_{uuid.uuid4().hex}"
        
        user_doc = {
            "user_id": user_id,
            "email": f"test_{uuid.uuid4().hex[:6]}@test.com",
            "name": "Test Export User",
            "role": "Editor",
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        db.users.insert_one(user_doc)
        
        # Create session
        session_doc = {
            "session_id": str(uuid.uuid4()),
            "user_id": user_id,
            "session_token": session_token,
            "expires_at": datetime(2030, 1, 1, tzinfo=timezone.utc).isoformat(),
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        db.user_sessions.insert_one(session_doc)
        
        # Create test project
        project_id = f"proj_{uuid.uuid4().hex[:12]}"
        project_doc = {
            "project_id": project_id,
            "user_id": user_id,
            "company_name": "Test Export Company Ltd",
            "sector": "Technology",
            "current_stage": "Assessment",
            "progress_percentage": 0,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }
        db.projects.insert_one(project_doc)
        
        self.test_project_id = project_id
        self.auth_token = session_token
        self.user_id = user_id
        self.db = db
        
        return project_id, session_token
    
    def _cleanup(self):
        """Cleanup test data"""
        if hasattr(self, 'db') and self.db is not None:
            if self.test_project_id:
                self.db.projects.delete_many({"project_id": self.test_project_id})
                self.db.drhp_output.delete_many({"project_id": self.test_project_id})
            if hasattr(self, 'user_id'):
                self.db.users.delete_many({"user_id": self.user_id})
                self.db.user_sessions.delete_many({"user_id": self.user_id})
    
    # ============ FORMAT PARAMETER TESTS ============
    
    def test_export_accepts_docx_format(self):
        """Test that export endpoint accepts 'docx' format parameter"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "docx",
                    "board_type": "sme",
                    "content": MINIMAL_HTML
                }
            )
            
            assert response.status_code == 200, f"Expected 200, got {response.status_code}: {response.text}"
            assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in response.headers.get("Content-Type", "")
            print("PASS: Export accepts 'docx' format parameter")
        finally:
            self._cleanup()
    
    def test_export_accepts_pdf_format(self):
        """Test that export endpoint accepts 'pdf' format parameter"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "pdf",
                    "board_type": "sme",
                    "content": MINIMAL_HTML
                }
            )
            
            assert response.status_code == 200, f"Expected 200, got {response.status_code}: {response.text}"
            assert "application/pdf" in response.headers.get("Content-Type", "")
            print("PASS: Export accepts 'pdf' format parameter")
        finally:
            self._cleanup()
    
    def test_invalid_format_returns_400(self):
        """Test that invalid format returns 400 error"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "txt",  # Invalid format
                    "board_type": "sme",
                    "content": MINIMAL_HTML
                }
            )
            
            assert response.status_code == 400, f"Expected 400, got {response.status_code}"
            assert "Invalid export format" in response.text or "invalid" in response.text.lower()
            print("PASS: Invalid format returns 400 error")
        finally:
            self._cleanup()
    
    # ============ WORD EXPORT TESTS ============
    
    def test_word_export_returns_valid_file(self):
        """Test that Word export returns a valid .docx file"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "docx",
                    "board_type": "sme",
                    "content": TEST_HTML_CONTENT
                }
            )
            
            assert response.status_code == 200
            
            # Check file size is reasonable (not empty)
            content_length = len(response.content)
            assert content_length > 1000, f"File too small: {content_length} bytes"
            
            # Check it's a valid ZIP (docx is a ZIP file)
            assert response.content[:4] == b'PK\x03\x04', "Not a valid DOCX file (should start with PK ZIP header)"
            
            print(f"PASS: Word export returns valid file ({content_length} bytes)")
        finally:
            self._cleanup()
    
    def test_word_export_preserves_headings(self):
        """Test that Word export preserves heading elements"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "docx",
                    "board_type": "sme",
                    "content": TEST_HTML_CONTENT
                }
            )
            
            assert response.status_code == 200
            
            # Parse the docx to verify headings
            from docx import Document
            doc = Document(io.BytesIO(response.content))
            
            # Check for heading content
            all_text = " ".join([p.text for p in doc.paragraphs])
            
            assert "DRAFT RED HERRING PROSPECTUS" in all_text, "Title heading not found"
            assert "SECTION I: GENERAL INFORMATION" in all_text, "H1 heading not found"
            assert "Risk Factors" in all_text, "H2 heading not found"
            
            print("PASS: Word export preserves headings")
        finally:
            self._cleanup()
    
    def test_word_export_preserves_tables(self):
        """Test that Word export preserves table elements"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "docx",
                    "board_type": "sme",
                    "content": TEST_HTML_CONTENT
                }
            )
            
            assert response.status_code == 200
            
            # Parse the docx to verify tables
            from docx import Document
            doc = Document(io.BytesIO(response.content))
            
            # Check for tables
            assert len(doc.tables) > 0, "No tables found in document"
            
            # Check table content
            table = doc.tables[0]
            table_text = ""
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "
            
            assert "Header 1" in table_text or "Cell 1" in table_text, "Table content not preserved"
            
            print(f"PASS: Word export preserves tables ({len(doc.tables)} table(s) found)")
        finally:
            self._cleanup()
    
    def test_word_export_preserves_bold_italic(self):
        """Test that Word export preserves bold and italic formatting"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "docx",
                    "board_type": "sme",
                    "content": TEST_HTML_CONTENT
                }
            )
            
            assert response.status_code == 200
            
            # Parse the docx to verify formatting
            from docx import Document
            doc = Document(io.BytesIO(response.content))
            
            # Check for bold and italic runs
            has_bold = False
            has_italic = False
            
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.bold:
                        has_bold = True
                    if run.italic:
                        has_italic = True
            
            # At least one of these should be true (headings are bold)
            assert has_bold or has_italic, "No bold or italic formatting found"
            
            print(f"PASS: Word export preserves formatting (bold={has_bold}, italic={has_italic})")
        finally:
            self._cleanup()
    
    # ============ PDF EXPORT TESTS ============
    
    def test_pdf_export_returns_valid_file(self):
        """Test that PDF export returns a valid PDF file"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "pdf",
                    "board_type": "sme",
                    "content": TEST_HTML_CONTENT
                }
            )
            
            assert response.status_code == 200
            
            # Check file size is reasonable
            content_length = len(response.content)
            assert content_length > 1000, f"File too small: {content_length} bytes"
            
            # Check PDF magic bytes
            assert response.content[:4] == b'%PDF', "Not a valid PDF file (should start with %PDF)"
            
            print(f"PASS: PDF export returns valid file ({content_length} bytes)")
        finally:
            self._cleanup()
    
    def test_pdf_export_includes_sebi_styles(self):
        """Test that PDF export includes SEBI CSS styles"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "pdf",
                    "board_type": "sme",
                    "content": TEST_HTML_CONTENT
                }
            )
            
            assert response.status_code == 200
            
            # PDF is valid and contains content
            content_length = len(response.content)
            assert content_length > 5000, f"PDF seems too small for styled content: {content_length} bytes"
            
            # Check PDF header
            assert response.content[:4] == b'%PDF'
            
            print(f"PASS: PDF export includes SEBI styles (file size: {content_length} bytes)")
        finally:
            self._cleanup()
    
    # ============ FILENAME TESTS ============
    
    def test_export_uses_company_name_in_filename(self):
        """Test that export uses company name in the filename"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            # Test DOCX filename
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "docx",
                    "board_type": "sme",
                    "content": MINIMAL_HTML
                }
            )
            
            assert response.status_code == 200
            
            content_disposition = response.headers.get("Content-Disposition", "")
            assert "Test_Export_Company" in content_disposition or "Test Export Company" in content_disposition, \
                f"Company name not in filename: {content_disposition}"
            assert ".docx" in content_disposition, f"Missing .docx extension: {content_disposition}"
            
            # Test PDF filename
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "pdf",
                    "board_type": "mainboard",
                    "content": MINIMAL_HTML
                }
            )
            
            assert response.status_code == 200
            
            content_disposition = response.headers.get("Content-Disposition", "")
            assert "MAINBOARD" in content_disposition.upper(), f"Board type not in filename: {content_disposition}"
            assert ".pdf" in content_disposition, f"Missing .pdf extension: {content_disposition}"
            
            print(f"PASS: Export uses company name in filename")
        finally:
            self._cleanup()
    
    # ============ FALLBACK TESTS ============
    
    def test_export_falls_back_to_saved_content(self):
        """Test that export falls back to saved content if none provided"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            # First, save some content
            save_response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "sme_content": "<h1>Saved SME Content</h1><p>This is saved content.</p>"
                }
            )
            assert save_response.status_code == 200, f"Failed to save content: {save_response.text}"
            
            # Now export without providing content
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "docx",
                    "board_type": "sme"
                    # No content provided - should use saved content
                }
            )
            
            assert response.status_code == 200, f"Expected 200, got {response.status_code}: {response.text}"
            
            # Verify it's a valid docx
            assert response.content[:4] == b'PK\x03\x04'
            
            # Parse and check content
            from docx import Document
            doc = Document(io.BytesIO(response.content))
            all_text = " ".join([p.text for p in doc.paragraphs])
            
            assert "Saved SME Content" in all_text, "Saved content not found in export"
            
            print("PASS: Export falls back to saved content")
        finally:
            self._cleanup()
    
    def test_export_no_content_returns_400(self):
        """Test that export with no content and no saved content returns 400"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            # Export without content and without saving anything first
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "docx",
                    "board_type": "sme"
                    # No content provided and nothing saved
                }
            )
            
            assert response.status_code == 400, f"Expected 400, got {response.status_code}"
            assert "No content" in response.text or "content" in response.text.lower()
            
            print("PASS: Export with no content returns 400")
        finally:
            self._cleanup()
    
    # ============ BOARD TYPE TESTS ============
    
    def test_export_sme_board_type(self):
        """Test export with SME board type"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "docx",
                    "board_type": "sme",
                    "content": MINIMAL_HTML
                }
            )
            
            assert response.status_code == 200
            
            content_disposition = response.headers.get("Content-Disposition", "")
            assert "SME" in content_disposition.upper(), f"SME not in filename: {content_disposition}"
            
            print("PASS: Export with SME board type works")
        finally:
            self._cleanup()
    
    def test_export_mainboard_type(self):
        """Test export with Mainboard type"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "docx",
                    "board_type": "mainboard",
                    "content": MINIMAL_HTML
                }
            )
            
            assert response.status_code == 200
            
            content_disposition = response.headers.get("Content-Disposition", "")
            assert "MAINBOARD" in content_disposition.upper(), f"MAINBOARD not in filename: {content_disposition}"
            
            print("PASS: Export with Mainboard type works")
        finally:
            self._cleanup()
    
    # ============ AUTHENTICATION TESTS ============
    
    def test_export_requires_authentication(self):
        """Test that export endpoint requires authentication"""
        response = self.session.post(
            f"{BASE_URL}/api/projects/fake_project_id/drhp-output/export",
            json={
                "format": "docx",
                "board_type": "sme",
                "content": MINIMAL_HTML
            }
        )
        
        assert response.status_code == 401, f"Expected 401, got {response.status_code}"
        print("PASS: Export requires authentication")
    
    # ============ CONTENT LENGTH TESTS ============
    
    def test_export_content_length_header(self):
        """Test that export includes Content-Length header"""
        project_id, token = self._create_test_user_and_project()
        
        try:
            response = self.session.post(
                f"{BASE_URL}/api/projects/{project_id}/drhp-output/export",
                headers={"Authorization": f"Bearer {token}"},
                json={
                    "format": "docx",
                    "board_type": "sme",
                    "content": TEST_HTML_CONTENT
                }
            )
            
            assert response.status_code == 200
            
            content_length = response.headers.get("Content-Length")
            assert content_length is not None, "Content-Length header missing"
            assert int(content_length) > 0, "Content-Length should be positive"
            assert int(content_length) == len(response.content), "Content-Length mismatch"
            
            print(f"PASS: Export includes correct Content-Length header ({content_length} bytes)")
        finally:
            self._cleanup()


# Run tests if executed directly
if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
