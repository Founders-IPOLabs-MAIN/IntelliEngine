"""
Test suite for DRHP Word Document Upload Feature
Tests the direct upload button functionality, file validation, and formatting preservation
"""
import pytest
import requests
import os
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', '').rstrip('/')

# Test session token - will be created in setup
SESSION_TOKEN = None
PROJECT_ID = None
USER_ID = None


@pytest.fixture(scope="module", autouse=True)
def setup_test_data():
    """Create test user, session, and project for testing"""
    global SESSION_TOKEN, PROJECT_ID, USER_ID
    
    import subprocess
    import time
    
    timestamp = int(time.time() * 1000)
    USER_ID = f"test-user-{timestamp}"
    SESSION_TOKEN = f"test_session_{timestamp}"
    PROJECT_ID = f"proj_test_{timestamp}"
    
    # Create test data in MongoDB
    mongo_script = f"""
    use('test_database');
    db.users.insertOne({{
      user_id: '{USER_ID}',
      email: 'test.user.{timestamp}@example.com',
      name: 'Test User',
      picture: 'https://via.placeholder.com/150',
      role: 'Editor',
      company_id: null,
      created_at: new Date().toISOString()
    }});
    db.user_sessions.insertOne({{
      session_id: 'sess_{timestamp}',
      user_id: '{USER_ID}',
      session_token: '{SESSION_TOKEN}',
      expires_at: new Date(Date.now() + 7*24*60*60*1000).toISOString(),
      created_at: new Date().toISOString()
    }});
    db.projects.insertOne({{
      project_id: '{PROJECT_ID}',
      user_id: '{USER_ID}',
      company_name: 'Test Corp',
      sector: 'Technology',
      current_stage: 'Assessment',
      progress_percentage: 0,
      created_at: new Date().toISOString(),
      updated_at: new Date().toISOString()
    }});
    """
    
    result = subprocess.run(
        ['mongosh', '--quiet', '--eval', mongo_script],
        capture_output=True,
        text=True
    )
    
    yield
    
    # Cleanup test data
    cleanup_script = f"""
    use('test_database');
    db.users.deleteMany({{user_id: '{USER_ID}'}});
    db.user_sessions.deleteMany({{session_token: '{SESSION_TOKEN}'}});
    db.projects.deleteMany({{project_id: '{PROJECT_ID}'}});
    db.drhp_output.deleteMany({{project_id: '{PROJECT_ID}'}});
    """
    subprocess.run(['mongosh', '--quiet', '--eval', cleanup_script], capture_output=True)


def create_test_docx(content_type="basic"):
    """Create a test .docx file with various formatting"""
    doc = Document()
    
    if content_type == "basic":
        doc.add_heading('Test Document', level=1)
        doc.add_paragraph('This is a test paragraph.')
    
    elif content_type == "formatted":
        # Title
        title = doc.add_heading('DRAFT RED HERRING PROSPECTUS', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Formatted text
        p = doc.add_paragraph()
        p.add_run('This is ').bold = False
        p.add_run('bold text').bold = True
        p.add_run(' and this is ')
        p.add_run('italic text').italic = True
        p.add_run(' and this is ')
        run = p.add_run('underlined text')
        run.underline = True
        
        # Heading
        doc.add_heading('Section 1: Risk Factors', level=2)
        doc.add_paragraph('Risk factors content here.')
    
    elif content_type == "with_table":
        doc.add_heading('Document with Table', level=1)
        
        # Add table
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Amount'
        row1 = table.rows[1].cells
        row1[0].text = 'Fresh Issue'
        row1[1].text = 'New shares'
        row1[2].text = 'Rs. 100 Cr'
        row2 = table.rows[2].cells
        row2[0].text = 'OFS'
        row2[1].text = 'Offer for Sale'
        row2[2].text = 'Rs. 50 Cr'
    
    elif content_type == "multi_page":
        # Create a multi-page document
        for i in range(10):
            doc.add_heading(f'Chapter {i+1}', level=1)
            for j in range(5):
                doc.add_paragraph(f'This is paragraph {j+1} of chapter {i+1}. ' * 10)
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


class TestAuthentication:
    """Test authentication for DRHP upload endpoints"""
    
    def test_auth_required_for_import(self):
        """Test that import endpoint requires authentication"""
        response = requests.post(
            f"{BASE_URL}/api/projects/test_project/drhp-output/import",
            files={'file': ('test.docx', b'fake content', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        )
        assert response.status_code == 401, f"Expected 401, got {response.status_code}"
        print("PASS: Import endpoint requires authentication")
    
    def test_valid_auth_accepted(self):
        """Test that valid session token is accepted"""
        response = requests.get(
            f"{BASE_URL}/api/auth/me",
            headers={"Authorization": f"Bearer {SESSION_TOKEN}"}
        )
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        data = response.json()
        assert data['user_id'] == USER_ID
        print("PASS: Valid authentication accepted")


class TestFileValidation:
    """Test file type and size validation"""
    
    def test_reject_txt_file(self):
        """Test that .txt files are rejected"""
        response = requests.post(
            f"{BASE_URL}/api/projects/{PROJECT_ID}/drhp-output/import?board_type=sme",
            headers={"Authorization": f"Bearer {SESSION_TOKEN}"},
            files={'file': ('test.txt', b'This is a text file', 'text/plain')}
        )
        assert response.status_code == 400, f"Expected 400, got {response.status_code}"
        assert "Only .docx files are supported" in response.json().get('detail', '')
        print("PASS: .txt files rejected correctly")
    
    def test_reject_pdf_file(self):
        """Test that .pdf files are rejected"""
        response = requests.post(
            f"{BASE_URL}/api/projects/{PROJECT_ID}/drhp-output/import?board_type=sme",
            headers={"Authorization": f"Bearer {SESSION_TOKEN}"},
            files={'file': ('test.pdf', b'%PDF-1.4 fake pdf', 'application/pdf')}
        )
        assert response.status_code == 400, f"Expected 400, got {response.status_code}"
        assert "Only .docx files are supported" in response.json().get('detail', '')
        print("PASS: .pdf files rejected correctly")
    
    def test_reject_doc_file(self):
        """Test that old .doc files are rejected"""
        response = requests.post(
            f"{BASE_URL}/api/projects/{PROJECT_ID}/drhp-output/import?board_type=sme",
            headers={"Authorization": f"Bearer {SESSION_TOKEN}"},
            files={'file': ('test.doc', b'fake doc content', 'application/msword')}
        )
        assert response.status_code == 400, f"Expected 400, got {response.status_code}"
        print("PASS: .doc files rejected correctly")
    
    def test_accept_docx_file(self):
        """Test that .docx files are accepted"""
        docx_buffer = create_test_docx("basic")
        response = requests.post(
            f"{BASE_URL}/api/projects/{PROJECT_ID}/drhp-output/import?board_type=sme",
            headers={"Authorization": f"Bearer {SESSION_TOKEN}"},
            files={'file': ('test.docx', docx_buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        )
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        data = response.json()
        assert data['success'] == True
        print("PASS: .docx files accepted correctly")


class TestFormattingPreservation:
    """Test that document formatting is preserved during import"""
    
    def test_bold_italic_underline_preserved(self):
        """Test that bold, italic, and underline formatting is preserved"""
        docx_buffer = create_test_docx("formatted")
        response = requests.post(
            f"{BASE_URL}/api/projects/{PROJECT_ID}/drhp-output/import?board_type=sme",
            headers={"Authorization": f"Bearer {SESSION_TOKEN}"},
            files={'file': ('formatted.docx', docx_buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        )
        assert response.status_code == 200
        data = response.json()
        html_content = data['html_content']
        
        # Check for formatting tags
        assert '<strong>' in html_content or '<b>' in html_content, "Bold formatting not preserved"
        assert '<em>' in html_content or '<i>' in html_content, "Italic formatting not preserved"
        assert '<u>' in html_content, "Underline formatting not preserved"
        print("PASS: Bold, italic, underline formatting preserved")
    
    def test_headings_preserved(self):
        """Test that headings are preserved"""
        docx_buffer = create_test_docx("formatted")
        response = requests.post(
            f"{BASE_URL}/api/projects/{PROJECT_ID}/drhp-output/import?board_type=sme",
            headers={"Authorization": f"Bearer {SESSION_TOKEN}"},
            files={'file': ('headings.docx', docx_buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        )
        assert response.status_code == 200
        data = response.json()
        html_content = data['html_content']
        
        # Check for heading tags
        assert '<h1>' in html_content, "H1 heading not preserved"
        assert '<h2>' in html_content, "H2 heading not preserved"
        print("PASS: Headings preserved")
    
    def test_tables_preserved(self):
        """Test that tables are preserved"""
        docx_buffer = create_test_docx("with_table")
        response = requests.post(
            f"{BASE_URL}/api/projects/{PROJECT_ID}/drhp-output/import?board_type=sme",
            headers={"Authorization": f"Bearer {SESSION_TOKEN}"},
            files={'file': ('table.docx', docx_buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        )
        assert response.status_code == 200
        data = response.json()
        html_content = data['html_content']
        
        # Check for table tags
        assert '<table' in html_content, "Table not preserved"
        assert '<tr>' in html_content, "Table rows not preserved"
        assert '<td>' in html_content, "Table cells not preserved"
        assert 'drhp-table' in html_content, "Table class not applied"
        print("PASS: Tables preserved with correct class")


class TestBoardTypeSelection:
    """Test SME and Mainboard board type selection"""
    
    def test_import_to_sme_board(self):
        """Test importing to SME board"""
        docx_buffer = create_test_docx("basic")
        response = requests.post(
            f"{BASE_URL}/api/projects/{PROJECT_ID}/drhp-output/import?board_type=sme",
            headers={"Authorization": f"Bearer {SESSION_TOKEN}"},
            files={'file': ('sme.docx', docx_buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        )
        assert response.status_code == 200
        data = response.json()
        assert data['board_type'] == 'sme'
        print("PASS: Import to SME board works")
    
    def test_import_to_mainboard(self):
        """Test importing to Mainboard"""
        docx_buffer = create_test_docx("basic")
        response = requests.post(
            f"{BASE_URL}/api/projects/{PROJECT_ID}/drhp-output/import?board_type=mainboard",
            headers={"Authorization": f"Bearer {SESSION_TOKEN}"},
            files={'file': ('mainboard.docx', docx_buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        )
        assert response.status_code == 200
        data = response.json()
        assert data['board_type'] == 'mainboard'
        print("PASS: Import to Mainboard works")


class TestResponseStructure:
    """Test API response structure"""
    
    def test_success_response_structure(self):
        """Test that success response has correct structure"""
        docx_buffer = create_test_docx("basic")
        response = requests.post(
            f"{BASE_URL}/api/projects/{PROJECT_ID}/drhp-output/import?board_type=sme",
            headers={"Authorization": f"Bearer {SESSION_TOKEN}"},
            files={'file': ('structure.docx', docx_buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        )
        assert response.status_code == 200
        data = response.json()
        
        # Check required fields
        assert 'success' in data, "Missing 'success' field"
        assert 'html_content' in data, "Missing 'html_content' field"
        assert 'board_type' in data, "Missing 'board_type' field"
        assert 'filename' in data, "Missing 'filename' field"
        assert 'file_size' in data, "Missing 'file_size' field"
        assert 'warnings' in data, "Missing 'warnings' field"
        assert 'warnings_count' in data, "Missing 'warnings_count' field"
        
        assert data['success'] == True
        assert isinstance(data['html_content'], str)
        assert isinstance(data['file_size'], int)
        assert isinstance(data['warnings'], list)
        print("PASS: Response structure is correct")


class TestMultiPageDocuments:
    """Test multi-page document support"""
    
    def test_multi_page_document_import(self):
        """Test that multi-page documents are imported correctly"""
        docx_buffer = create_test_docx("multi_page")
        response = requests.post(
            f"{BASE_URL}/api/projects/{PROJECT_ID}/drhp-output/import?board_type=sme",
            headers={"Authorization": f"Bearer {SESSION_TOKEN}"},
            files={'file': ('multipage.docx', docx_buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        )
        assert response.status_code == 200
        data = response.json()
        
        # Check that all chapters are present
        html_content = data['html_content']
        for i in range(1, 11):
            assert f'Chapter {i}' in html_content, f"Chapter {i} not found in imported content"
        
        print("PASS: Multi-page document imported with all content")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
