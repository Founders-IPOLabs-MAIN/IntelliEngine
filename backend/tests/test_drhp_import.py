"""
Test suite for DRHP Word Document Import functionality
Tests the POST /api/projects/{project_id}/drhp-output/import endpoint
"""
import pytest
import requests
import os
import subprocess
import time
import io

# Get BASE_URL from environment
BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', '').rstrip('/')


@pytest.fixture(scope="module")
def test_data():
    """Create test user, session, and project via MongoDB"""
    timestamp = int(time.time() * 1000)
    
    # Create test data via mongosh
    result = subprocess.run([
        'mongosh', '--quiet', '--eval', f'''
        use('test_database');
        var userId = 'test-user-import-pytest-{timestamp}';
        var sessionToken = 'test_session_import_pytest_{timestamp}';
        var projectId = 'proj_import_pytest_{timestamp}';
        
        db.users.insertOne({{
          user_id: userId,
          email: 'test.import.pytest.{timestamp}@example.com',
          name: 'Import Pytest User',
          picture: 'https://via.placeholder.com/150',
          role: 'Editor',
          company_id: null,
          created_at: new Date().toISOString()
        }});
        
        db.user_sessions.insertOne({{
          session_id: 'sess_import_pytest_{timestamp}',
          user_id: userId,
          session_token: sessionToken,
          expires_at: new Date(Date.now() + 7*24*60*60*1000).toISOString(),
          created_at: new Date().toISOString()
        }});
        
        db.projects.insertOne({{
          project_id: projectId,
          user_id: userId,
          company_name: 'Import Pytest Corp',
          sector: 'Technology',
          current_stage: 'Assessment',
          progress_percentage: 0,
          created_at: new Date().toISOString(),
          updated_at: new Date().toISOString()
        }});
        
        print('SESSION_TOKEN=' + sessionToken);
        print('USER_ID=' + userId);
        print('PROJECT_ID=' + projectId);
        '''
    ], capture_output=True, text=True)
    
    # Parse output
    data = {}
    for line in result.stdout.split('\n'):
        if line.startswith('SESSION_TOKEN='):
            data['session_token'] = line.split('=')[1]
        elif line.startswith('USER_ID='):
            data['user_id'] = line.split('=')[1]
        elif line.startswith('PROJECT_ID='):
            data['project_id'] = line.split('=')[1]
    
    yield data
    
    # Cleanup after tests
    subprocess.run([
        'mongosh', '--quiet', '--eval', f'''
        use('test_database');
        db.users.deleteMany({{user_id: /test-user-import-pytest-{timestamp}/}});
        db.user_sessions.deleteMany({{session_token: /test_session_import_pytest_{timestamp}/}});
        db.projects.deleteMany({{project_id: /proj_import_pytest_{timestamp}/}});
        db.drhp_output.deleteMany({{project_id: /proj_import_pytest_{timestamp}/}});
        '''
    ], capture_output=True, text=True)


@pytest.fixture
def api_client(test_data):
    """Create requests session with auth header"""
    session = requests.Session()
    session.headers.update({
        "Authorization": f"Bearer {test_data['session_token']}"
    })
    return session


@pytest.fixture
def sample_docx():
    """Create a sample .docx file for testing"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('Test DRHP Document', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add formatted text
    p = doc.add_paragraph()
    p.add_run('Bold text').bold = True
    p.add_run(' and ')
    p.add_run('italic text').italic = True
    
    # Add heading
    doc.add_heading('Section 1', level=2)
    doc.add_paragraph('Test paragraph content.')
    
    # Add table
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Header 1'
    table.rows[0].cells[1].text = 'Header 2'
    table.rows[1].cells[0].text = 'Data 1'
    table.rows[1].cells[1].text = 'Data 2'
    
    # Add list
    doc.add_paragraph('List item 1', style='List Bullet')
    doc.add_paragraph('List item 2', style='List Bullet')
    
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============ TESTS ============

class TestDRHPImportAPI:
    """Test DRHP Word Import API endpoint"""
    
    def test_auth_works(self, api_client, test_data):
        """Test that authentication is working"""
        response = api_client.get(f"{BASE_URL}/api/auth/me")
        assert response.status_code == 200, f"Auth failed: {response.text}"
        data = response.json()
        assert "user_id" in data
        assert data["user_id"] == test_data['user_id']
        print(f"Auth test passed - User: {data['name']}")
    
    def test_import_docx_sme(self, api_client, test_data, sample_docx):
        """Test importing .docx file to SME board"""
        files = {
            'file': ('test_document.docx', sample_docx, 
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 200, f"Import failed: {response.text}"
        data = response.json()
        
        # Verify response structure
        assert data["success"] == True
        assert "html_content" in data
        assert data["board_type"] == "sme"
        assert data["filename"] == "test_document.docx"
        assert "file_size" in data
        assert "warnings" in data
        assert "warnings_count" in data
        
        # Verify HTML content contains expected elements
        html = data["html_content"]
        assert "<h1" in html  # Title (may have class attribute)
        assert "<strong>" in html  # Bold text
        assert "<em>" in html  # Italic text
        assert "<h2" in html  # Section heading (may have class attribute)
        assert "<table" in html  # Table
        # List items may be rendered as <li> or <p class="drhp-list-bullet">
        assert "List item" in html  # List content preserved
        
        print(f"Import SME test passed - File size: {data['file_size']} bytes")
    
    def test_import_docx_mainboard(self, api_client, test_data, sample_docx):
        """Test importing .docx file to Mainboard"""
        files = {
            'file': ('test_mainboard.docx', sample_docx,
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=mainboard",
            files=files
        )
        
        assert response.status_code == 200, f"Import failed: {response.text}"
        data = response.json()
        assert data["success"] == True
        assert data["board_type"] == "mainboard"
        print("Import Mainboard test passed")
    
    def test_content_saved_to_database(self, api_client, test_data, sample_docx):
        """Test that imported content is saved to database"""
        # First import a document
        files = {
            'file': ('test_save.docx', sample_docx,
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        
        import_response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        assert import_response.status_code == 200, f"Import failed: {import_response.text}"
        
        # Then verify it was saved
        get_response = api_client.get(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output"
        )
        assert get_response.status_code == 200
        data = get_response.json()
        
        # Verify content was saved
        assert "sme_content" in data
        assert len(data["sme_content"]) > 0
        
        # Verify import info was saved
        assert "sme_import_info" in data
        assert data["sme_import_info"]["filename"] == "test_save.docx"
        assert "imported_at" in data["sme_import_info"]
        assert "file_size" in data["sme_import_info"]
        
        print(f"Content saved test passed - Import info: {data['sme_import_info']}")
    
    def test_reject_non_docx_file(self, api_client, test_data):
        """Test that non-.docx files are rejected"""
        # Create a fake text file
        fake_file = io.BytesIO(b"This is not a docx file")
        
        files = {
            'file': ('test.txt', fake_file, 'text/plain')
        }
        
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 400
        data = response.json()
        assert "Only .docx files are supported" in data["detail"]
        print("Non-docx rejection test passed")
    
    def test_reject_pdf_file(self, api_client, test_data):
        """Test that PDF files are rejected"""
        fake_pdf = io.BytesIO(b"%PDF-1.4 fake pdf content")
        
        files = {
            'file': ('document.pdf', fake_pdf, 'application/pdf')
        }
        
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 400
        data = response.json()
        assert "Only .docx files are supported" in data["detail"]
        print("PDF rejection test passed")
    
    def test_tables_preserved(self, api_client, test_data, sample_docx):
        """Test that tables are preserved in imported content"""
        files = {
            'file': ('test_tables.docx', sample_docx,
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 200, f"Import failed: {response.text}"
        data = response.json()
        html = data["html_content"]
        
        # Verify table structure
        assert '<table class="drhp-table">' in html
        assert '<tr>' in html
        assert '<td>' in html
        assert 'Header 1' in html
        assert 'Data 1' in html
        
        print("Tables preserved test passed")
    
    def test_text_formatting_preserved(self, api_client, test_data, sample_docx):
        """Test that text formatting (bold, italic) is preserved"""
        files = {
            'file': ('test_formatting.docx', sample_docx,
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 200, f"Import failed: {response.text}"
        data = response.json()
        html = data["html_content"]
        
        # Verify formatting tags
        assert '<strong>' in html  # Bold
        assert '<em>' in html  # Italic
        
        print("Text formatting preserved test passed")
    
    def test_headings_preserved(self, api_client, test_data, sample_docx):
        """Test that headings are preserved"""
        files = {
            'file': ('test_headings.docx', sample_docx,
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 200, f"Import failed: {response.text}"
        data = response.json()
        html = data["html_content"]
        
        # Verify heading tags (may have class attributes)
        assert '<h1' in html
        assert '<h2' in html
        
        print("Headings preserved test passed")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
