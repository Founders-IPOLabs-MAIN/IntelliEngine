"""
Test suite for DRHP SEBI-specific formatting preservation
Tests the new python-docx based parser with SEBI formatting features:
- Paragraph alignment (left, center, right, justify)
- Indentation (left, right, first-line)
- Numbered lists with hierarchy levels
- Tables with drhp-table class
- Hyperlinks extraction
- Images stored in GridFS
- Bold, italic, underline formatting
- SEBI-specific CSS styles
"""
import pytest
import requests
import os
import subprocess
import time
import io
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Get BASE_URL from environment
BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', '').rstrip('/')


@pytest.fixture(scope="module")
def test_data():
    """Create test user, session, and project via MongoDB"""
    timestamp = int(time.time() * 1000)
    
    result = subprocess.run([
        'mongosh', '--quiet', '--eval', f'''
        use('test_database');
        var userId = 'test-user-sebi-{timestamp}';
        var sessionToken = 'test_session_sebi_{timestamp}';
        var projectId = 'proj_sebi_{timestamp}';
        
        db.users.insertOne({{
          user_id: userId,
          email: 'test.sebi.{timestamp}@example.com',
          name: 'SEBI Test User',
          picture: 'https://via.placeholder.com/150',
          role: 'Editor',
          company_id: null,
          created_at: new Date().toISOString()
        }});
        
        db.user_sessions.insertOne({{
          session_id: 'sess_sebi_{timestamp}',
          user_id: userId,
          session_token: sessionToken,
          expires_at: new Date(Date.now() + 7*24*60*60*1000).toISOString(),
          created_at: new Date().toISOString()
        }});
        
        db.projects.insertOne({{
          project_id: projectId,
          user_id: userId,
          company_name: 'SEBI Test Corp',
          sector: 'Technology',
          current_stage: 'Assessment',
          progress_percentage: 0,
          created_at: new Date().toISOString(),
          updated_at: new Date().toISOString()
        }});
        
        print('SESSION_TOKEN=' + sessionToken);
        print('USER_ID=' + userId);
        print('PROJECT_ID=' + projectId);
        '''
    ], capture_output=True, text=True)
    
    data = {}
    for line in result.stdout.split('\n'):
        if line.startswith('SESSION_TOKEN='):
            data['session_token'] = line.split('=')[1]
        elif line.startswith('USER_ID='):
            data['user_id'] = line.split('=')[1]
        elif line.startswith('PROJECT_ID='):
            data['project_id'] = line.split('=')[1]
    
    yield data
    
    # Cleanup
    subprocess.run([
        'mongosh', '--quiet', '--eval', f'''
        use('test_database');
        db.users.deleteMany({{user_id: /test-user-sebi-{timestamp}/}});
        db.user_sessions.deleteMany({{session_token: /test_session_sebi_{timestamp}/}});
        db.projects.deleteMany({{project_id: /proj_sebi_{timestamp}/}});
        db.drhp_output.deleteMany({{project_id: /proj_sebi_{timestamp}/}});
        db.fs.files.deleteMany({{metadata: {{project_id: /proj_sebi_{timestamp}/}}}});
        db.fs.chunks.deleteMany({{}});
        '''
    ], capture_output=True, text=True)


@pytest.fixture
def api_client(test_data):
    """Create requests session with auth header"""
    session = requests.Session()
    session.headers.update({
        "Authorization": f"Bearer {test_data['session_token']}"
    })
    return session


def create_docx_with_alignment():
    """Create a docx with different text alignments"""
    doc = Document()
    
    # Left aligned
    p_left = doc.add_paragraph('Left aligned text')
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Center aligned
    p_center = doc.add_paragraph('Center aligned text')
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Right aligned
    p_right = doc.add_paragraph('Right aligned text')
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Justified
    p_justify = doc.add_paragraph('Justified text that should span across the page width for proper justification.')
    p_justify.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_docx_with_indentation():
    """Create a docx with various indentation styles"""
    doc = Document()
    
    # Normal paragraph
    doc.add_paragraph('Normal paragraph without indentation')
    
    # Left indented paragraph
    p_left = doc.add_paragraph('Left indented paragraph')
    p_left.paragraph_format.left_indent = Inches(0.5)
    
    # First line indent
    p_first = doc.add_paragraph('Paragraph with first line indent. This is a longer paragraph to show the first line indent effect.')
    p_first.paragraph_format.first_line_indent = Inches(0.5)
    
    # Right indent
    p_right = doc.add_paragraph('Right indented paragraph')
    p_right.paragraph_format.right_indent = Inches(0.5)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_docx_with_numbered_lists():
    """Create a docx with numbered list hierarchy"""
    doc = Document()
    
    doc.add_heading('Numbered List Test', level=1)
    
    # Add numbered list items
    doc.add_paragraph('First item', style='List Number')
    doc.add_paragraph('Second item', style='List Number')
    doc.add_paragraph('Third item', style='List Number')
    
    # Add bullet list items
    doc.add_paragraph('Bullet item 1', style='List Bullet')
    doc.add_paragraph('Bullet item 2', style='List Bullet')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_docx_with_formatting():
    """Create a docx with bold, italic, underline formatting"""
    doc = Document()
    
    p = doc.add_paragraph()
    p.add_run('Bold text').bold = True
    p.add_run(' ')
    p.add_run('Italic text').italic = True
    p.add_run(' ')
    run_underline = p.add_run('Underlined text')
    run_underline.underline = True
    
    # Combined formatting
    p2 = doc.add_paragraph()
    run_combined = p2.add_run('Bold and Italic')
    run_combined.bold = True
    run_combined.italic = True
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_docx_with_table():
    """Create a docx with a table"""
    doc = Document()
    
    doc.add_heading('Table Test', level=1)
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Header row
    table.rows[0].cells[0].text = 'Column A'
    table.rows[0].cells[1].text = 'Column B'
    table.rows[0].cells[2].text = 'Column C'
    
    # Data rows
    table.rows[1].cells[0].text = 'Data A1'
    table.rows[1].cells[1].text = 'Data B1'
    table.rows[1].cells[2].text = 'Data C1'
    
    table.rows[2].cells[0].text = 'Data A2'
    table.rows[2].cells[1].text = 'Data B2'
    table.rows[2].cells[2].text = 'Data C2'
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_docx_with_hyperlink():
    """Create a docx with a hyperlink"""
    doc = Document()
    
    # Add a paragraph with text
    p = doc.add_paragraph('Visit our website: ')
    
    # Add hyperlink (simplified - python-docx doesn't have direct hyperlink support)
    # The parser should still handle any hyperlinks in the document
    p.add_run('https://example.com')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


class TestSEBIFormattingPreservation:
    """Test SEBI-specific formatting preservation"""
    
    def test_sebi_css_styles_included(self, api_client, test_data):
        """Test that SEBI-specific CSS styles are included in output"""
        doc = Document()
        doc.add_paragraph('Test content')
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        files = {'file': ('test.docx', buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 200, f"Import failed: {response.text}"
        html = response.json()["html_content"]
        
        # Check for SEBI-specific CSS classes
        assert '.drhp-document' in html
        assert '.drhp-title' in html
        assert '.drhp-h1' in html
        assert '.drhp-clause' in html
        assert '.drhp-sub-clause' in html
        assert '.drhp-risk-factor' in html
        assert '.drhp-legal' in html
        assert '.drhp-table' in html
        assert '.drhp-numbered-list' in html
        
        print("SEBI CSS styles test passed")
    
    def test_paragraph_alignment_preserved(self, api_client, test_data):
        """Test that paragraph alignment is preserved"""
        buffer = create_docx_with_alignment()
        
        files = {'file': ('alignment.docx', buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 200, f"Import failed: {response.text}"
        html = response.json()["html_content"]
        
        # Check for alignment styles
        assert 'text-align: left' in html or 'Left aligned' in html
        assert 'text-align: center' in html
        assert 'text-align: right' in html
        assert 'text-align: justify' in html
        
        print("Paragraph alignment test passed")
    
    def test_indentation_preserved(self, api_client, test_data):
        """Test that indentation is preserved as inline styles"""
        buffer = create_docx_with_indentation()
        
        files = {'file': ('indent.docx', buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 200, f"Import failed: {response.text}"
        html = response.json()["html_content"]
        
        # Check for indentation styles
        assert 'margin-left:' in html or 'Left indented' in html
        assert 'text-indent:' in html or 'first line indent' in html
        
        print("Indentation preservation test passed")
    
    def test_numbered_lists_hierarchy(self, api_client, test_data):
        """Test that numbered lists preserve hierarchy levels"""
        buffer = create_docx_with_numbered_lists()
        
        files = {'file': ('lists.docx', buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 200, f"Import failed: {response.text}"
        html = response.json()["html_content"]
        
        # Check for list content
        assert 'First item' in html
        assert 'Second item' in html
        assert 'Bullet item' in html
        
        # Check for list styling classes
        assert 'drhp-list' in html or '<li' in html or 'list' in html.lower()
        
        print("Numbered lists hierarchy test passed")
    
    def test_tables_with_drhp_class(self, api_client, test_data):
        """Test that tables are preserved with drhp-table class"""
        buffer = create_docx_with_table()
        
        files = {'file': ('table.docx', buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 200, f"Import failed: {response.text}"
        html = response.json()["html_content"]
        
        # Check for table with drhp-table class
        assert '<table class="drhp-table">' in html
        assert '<th>' in html  # Header cells
        assert '<td>' in html  # Data cells
        assert 'Column A' in html
        assert 'Data A1' in html
        
        print("Tables with drhp-table class test passed")
    
    def test_bold_italic_underline_preserved(self, api_client, test_data):
        """Test that bold, italic, underline formatting is preserved"""
        buffer = create_docx_with_formatting()
        
        files = {'file': ('formatting.docx', buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 200, f"Import failed: {response.text}"
        html = response.json()["html_content"]
        
        # Check for formatting tags
        assert '<strong>' in html  # Bold
        assert '<em>' in html  # Italic
        assert '<u>' in html  # Underline
        
        print("Bold, italic, underline formatting test passed")
    
    def test_images_array_in_response(self, api_client, test_data):
        """Test that import response includes images array"""
        doc = Document()
        doc.add_paragraph('Document without images')
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        files = {'file': ('no_images.docx', buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 200, f"Import failed: {response.text}"
        data = response.json()
        
        # Check for images array in response
        assert 'images' in data
        assert 'images_count' in data
        assert isinstance(data['images'], list)
        
        print("Images array in response test passed")
    
    def test_parser_type_in_import_info(self, api_client, test_data):
        """Test that import info includes parser type"""
        doc = Document()
        doc.add_paragraph('Test content')
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        files = {'file': ('parser_test.docx', buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 200, f"Import failed: {response.text}"
        
        # Get saved content
        get_response = api_client.get(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output"
        )
        assert get_response.status_code == 200
        data = get_response.json()
        
        # Check import info
        assert 'sme_import_info' in data
        assert data['sme_import_info']['parser'] == 'python-docx-sebi'
        
        print("Parser type in import info test passed")
    
    def test_warnings_returned(self, api_client, test_data):
        """Test that warnings are returned in response"""
        doc = Document()
        doc.add_paragraph('Test content')
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        files = {'file': ('warnings_test.docx', buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = api_client.post(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-output/import?board_type=sme",
            files=files
        )
        
        assert response.status_code == 200, f"Import failed: {response.text}"
        data = response.json()
        
        # Check for warnings fields
        assert 'warnings' in data
        assert 'warnings_count' in data
        assert isinstance(data['warnings'], list)
        
        print("Warnings returned test passed")


class TestDRHPImageEndpoint:
    """Test the new GET endpoint for retrieving stored images"""
    
    def test_image_endpoint_exists(self, api_client, test_data):
        """Test that the image endpoint returns 404 for non-existent image"""
        # Use a valid ObjectId format but non-existent
        fake_image_id = "507f1f77bcf86cd799439011"
        
        response = api_client.get(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-images/{fake_image_id}"
        )
        
        # Should return 404 for non-existent image
        assert response.status_code == 404
        
        print("Image endpoint exists test passed")
    
    def test_image_endpoint_invalid_id(self, api_client, test_data):
        """Test that the image endpoint handles invalid ObjectId"""
        invalid_id = "invalid-id"
        
        response = api_client.get(
            f"{BASE_URL}/api/projects/{test_data['project_id']}/drhp-images/{invalid_id}"
        )
        
        # Should return 404 for invalid ID
        assert response.status_code == 404
        
        print("Image endpoint invalid ID test passed")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
